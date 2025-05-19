import re
import docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree
from latex2mathml.converter import convert as _latex_to_mathml
from markdowntodocx.markdownconverter import convertMarkdownInFile

mml2omml_string = '''<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
  xmlns:xs="http://www.w3.org/2001/XMLSchema"
  xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
  xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
  xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
  xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
  xmlns:w10="urn:schemas-microsoft-com:office:word"
  xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
  xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"
  xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"
  xmlns:v="urn:schemas-microsoft-com:vml"
  xmlns:o="urn:schemas-microsoft-com:office:office"
  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
  xmlns:mml="http://www.w3.org/1998/Math/MathML"
  xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
  xmlns:ct="http://schemas.openxmlformats.org/package/2006/content-types"
  xmlns:docx2hub="http://transpect.io/docx2hub"
  xmlns:rel="http://schemas.openxmlformats.org/package/2006/relationships"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:c="http://www.w3.org/ns/xproc-step"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:omml="http://schemas.openxmlformats.org/officeDocument/2006/math"
  exclude-result-prefixes="xs rel docx2hub c"
  version="2.0">

  <xsl:template match="/">
    <xsl:apply-templates/>
  </xsl:template>

  <xsl:template match="mml:math">
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <xsl:apply-templates/>
    </m:oMath>
  </xsl:template>

  <xsl:template match="mml:mi | mml:mo | mml:mn">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>
        <xsl:value-of select="."/>
      </m:t>
    </m:r>
  </xsl:template>

  <!-- Superscript -->
  <xsl:template match="mml:msup">
    <m:sSup xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:e>
        <xsl:apply-templates select="mml:*[1]"/>
      </m:e>
      <m:sup>
        <xsl:apply-templates select="mml:*[2]"/>
      </m:sup>
    </m:sSup>
  </xsl:template>

  <!-- Subscript -->
  <xsl:template match="mml:msub">
    <m:sSub xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:e>
        <xsl:apply-templates select="mml:*[1]"/>
      </m:e>
      <m:sub>
        <xsl:apply-templates select="mml:*[2]"/>
      </m:sub>
    </m:sSub>
  </xsl:template>

  <!-- Subscript and Superscript -->
  <xsl:template match="mml:msubsup">
    <m:sSubSup xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:e>
        <xsl:apply-templates select="mml:*[1]"/>
      </m:e>
      <m:sub>
        <xsl:apply-templates select="mml:*[2]"/>
      </m:sub>
      <m:sup>
        <xsl:apply-templates select="mml:*[3]"/>
      </m:sup>
    </m:sSubSup>
  </xsl:template>

  <!-- Fraction -->
  <xsl:template match="mml:mfrac">
    <m:f xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:num>
        <xsl:apply-templates select="mml:*[1]"/>
      </m:num>
      <m:den>
        <xsl:apply-templates select="mml:*[2]"/>
      </m:den>
    </m:f>
  </xsl:template>

  <!-- Square Root -->
  <xsl:template match="mml:msqrt">
    <m:rad xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:deg/>
      <m:e>
        <xsl:apply-templates/>
      </m:e>
    </m:rad>
  </xsl:template>

  <!-- nth Root -->
  <xsl:template match="mml:mroot">
    <m:rad xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:deg>
        <xsl:apply-templates select="mml:*[2]"/>
      </m:deg>
      <m:e>
        <xsl:apply-templates select="mml:*[1]"/>
      </m:e>
    </m:rad>
  </xsl:template>

  <!-- Handling Delimiters as Regular Text -->
  <xsl:template match="mml:mo[.='(' or .=')' or .='[' or .=']' or .='{' or .='}']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>
        <xsl:value-of select="."/>
      </m:t>
    </m:r>
  </xsl:template>

  <!-- Handling Text within equations -->
  <xsl:template match="mml:mtext">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>
        <xsl:value-of select="."/>
      </m:t>
    </m:r>
  </xsl:template>

  <!-- Summation -->
  <xsl:template match="mml:mo[. = '∑']">
    <m:nary xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:naryPr>
        <m:chr>∑</m:chr>
      </m:naryPr>
      <xsl:apply-templates/>
    </m:nary>
  </xsl:template>

  <!-- Product -->
  <xsl:template match="mml:mo[. = '∏']">
    <m:nary xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:naryPr>
        <m:chr>∏</m:chr>
      </m:naryPr>
      <xsl:apply-templates/>
    </m:nary>
  </xsl:template>

  <!-- Integral -->
  <xsl:template match="mml:mo[. = '∫']">
    <m:nary xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:naryPr>
        <m:chr>∫</m:chr>
      </m:naryPr>
      <xsl:apply-templates/>
    </m:nary>
  </xsl:template>

  <!-- Greek Letters -->
  <xsl:template match="mml:mi[. = 'α']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>α</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mi[. = 'β']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>β</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mi[. = 'γ']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>γ</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mi[. = 'π']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>π</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mi[. = 'Σ']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>Σ</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mi[. = 'θ']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>θ</m:t>
    </m:r>
  </xsl:template>

  <!-- Handling General Mathematical Symbols -->
  <xsl:template match="mml:mo[. = '×']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>×</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mo[. = '÷']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>÷</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mo[. = '±']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>±</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mo[. = '∪']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>∪</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mo[. = '∩']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>∩</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mo[. = '≤']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>≤</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mo[. = '≥']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>≥</m:t>
    </m:r>
  </xsl:template>

  <xsl:template match="mml:mo[. = '≠']">
    <m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:t>≠</m:t>
    </m:r>
  </xsl:template>

</xsl:stylesheet>'''

# Parse the XSLT string
xslt = etree.XML(mml2omml_string.encode('utf-8'))
transform = etree.XSLT(xslt)

def latex_to_mathml(latex):
    mathml = _latex_to_mathml(latex)
    return mathml

def latex(doc, equation, centered=False, paragraph=None):
    '''
    Example equation=r'\sin^2(x)+\cos^2(x) - \frac{1}{2} \neq 1'
    '''
    if paragraph is None:
        paragraph = doc.add_paragraph()
    if centered:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Convert LaTeX to MathML
    mathml = latex_to_mathml(equation)

    # Transform MathML to OMML
    mathml_tree = etree.fromstring(mathml)
    omml_tree = transform(mathml_tree)

    # Create a new paragraph and append the OMML content
    paragraph._element.append(omml_tree.getroot())

    return paragraph

def docx_render_md(input_path, output_path):
    # Takes a docx file with markdown content and renders it to a formatted docx file
    convertMarkdownInFile(input_path, output_path, {"Code Car":"CodeStyle"})

def preprocess_images(markdown_text):
    """Convert HTML img tags to markdown image syntax."""
    def convert_github_url(url):
        # Convert GitHub web interface URLs to raw.githubusercontent.com URLs
        if 'github.com' in url and 'blob' in url:
            # Remove ?raw=true if present
            url = url.split('?')[0]
            # Convert to raw.githubusercontent.com URL
            url = url.replace('github.com', 'raw.githubusercontent.com')
            url = url.replace('/blob/', '/')
        return url

    def replace_img(match):
        src = convert_github_url(match.group(1))
        alt = match.group(2)
        return f"![{alt}]({src})"
    
    # Match img tags with src and alt attributes
    img_pattern = r'<img\s+src="([^"]+)"\s+alt="([^"]+)"[^>]*>'
    return re.sub(img_pattern, replace_img, markdown_text)

def preprocess_videos(markdown_text):
    """Convert video links and iframe embeds to their URLs."""
    def replace_video(match):
        src_url = match.group(1)
        return src_url
    
    # First handle iframe embeds, including those nested in divs
    iframe_pattern = r'<div[^>]*>.*?<iframe[^>]*src="([^"]+)"[^>]*>.*?</div>'
    markdown_text = re.sub(iframe_pattern, replace_video, markdown_text, flags=re.DOTALL)
    
    # Then handle direct video links in alert blocks
    alert_pattern = r'<div[^>]*class="alert[^"]*"[^>]*>.*?<a[^>]*href="([^"]+)"[^>]*>.*?</div>'
    markdown_text = re.sub(alert_pattern, replace_video, markdown_text, flags=re.DOTALL)
    
    return markdown_text

def preprocess_markdown(markdown_text):
    """Apply all preprocessing steps to the markdown text."""
    # Process videos first (before removing divs)
    markdown_text = preprocess_videos(markdown_text)
    
    # Remove any remaining divs
    markdown_text = re.sub(r'<div[^>]*>.*?</div>', '', markdown_text, flags=re.DOTALL)
    
    # Process images
    markdown_text = preprocess_images(markdown_text)
    
    return markdown_text

def markdownToWordFromString(string, outfile):
    # Preprocess the markdown string
    string = preprocess_markdown(string)
    
    document = docx.Document()
    
    # Add the required Cell style
    styles = document.styles
    if 'Cell' not in styles:
        cell_style = styles.add_style('Cell', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        cell_style.base_style = styles['Normal']
    
    paragraphs = string.replace("\r","").split("\n")
    current_paragraph = None
    
    for para in paragraphs:
        if not para.strip():
            continue
            
        # Check if this is a bullet point
        is_bullet = para.strip().startswith('- ')
        if is_bullet:
            para = para.strip()[2:]  # Remove the bullet point marker
            
        # Check if this is a blockquote and count the nesting level
        blockquote_level = 0
        if para.strip().startswith('>'):
            # Count the number of > symbols at the start of the line
            match = re.match(r'^(\s*>\s*)+', para)
            if match:
                # Count the number of > symbols in the matched string
                blockquote_level = match.group(0).count('>')
                # Remove all > symbols and their surrounding spaces
                para = para[match.end():].strip()
            
        # First handle display equations ($$)
        if '$$' in para:
            # Split the paragraph by $$ delimiters
            parts = re.split(r'(\$\$.*?\$\$)', para)
            
            for part in parts:
                if part.startswith('$$') and part.endswith('$$'):
                    # Extract the equation without the $$ delimiters
                    equation = part[2:-2]
                    # Use latex() function to add the equation
                    latex(document, equation, centered=True)
                else:
                    # Process inline equations in the text part
                    if part.strip():
                        process_inline_equations(document, part, is_bullet, blockquote_level)
        else:
            # Process inline equations in regular paragraphs
            process_inline_equations(document, para, is_bullet, blockquote_level)
    
    document.save(outfile)
    return convertMarkdownInFile(outfile, outfile)

def process_inline_equations(document, text, is_bullet=False, blockquote_level=0):
    if not text.strip():
        return
        
    # Split the text by $ delimiters, but don't capture the $ in the split
    parts = re.split(r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)', text)
    
    # Create a new paragraph for the text
    paragraph = document.add_paragraph()
    
    # Apply appropriate style
    if is_bullet:
        paragraph.style = 'List Bullet'
    elif blockquote_level > 0:
        paragraph.style = 'Quote'
        # Add indentation based on blockquote level
        paragraph.paragraph_format.left_indent = Inches(0.5 * blockquote_level)
    
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Even indices are regular text
            if part.strip():
                paragraph.add_run(part)
        else:
            # Odd indices are equations
            if part.strip():
                # Use latex() function to add the equation inline
                latex(document, part, centered=False, paragraph=paragraph)